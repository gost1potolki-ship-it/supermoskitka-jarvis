import json
from pathlib import Path
from statistics import mean

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DOCX = ROOT / "dealer-pleated-mesh-price-per-m2.docx"
SUMMARY_JSON = ROOT / "scripts" / "output" / "dealer-plisse-summary.json"
RANGE_JSON = ROOT / "scripts" / "output" / "dealer-plisse-range-summary.json"
CHECKS_JSON = ROOT / "scripts" / "output" / "dealer-plisse-size-checks.json"


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def money(value: float) -> str:
    return f"{value:,.2f}".replace(",", " ")


def percent(value: float) -> str:
    return f"{value:.2f}%"


def build_doc() -> None:
    summary_rows = load_json(SUMMARY_JSON)
    range_rows = load_json(RANGE_JSON)
    check_rows = load_json(CHECKS_JSON)

    all_spreads = [row["spreadPct"] for row in summary_rows]
    avg_spread = mean(all_spreads)
    min_spread = min(all_spreads)
    max_spread = max(all_spreads)
    ranges_required_count = sum(
        1 for row in summary_rows if row["recommendation"] == "лучше использовать диапазоны"
    )

    document = Document()
    document.add_heading("Стоимость 1 м² сеток плиссе для дилеров", level=1)

    document.add_paragraph(
        "Расчёт выполнен на базе действующего калькулятора проекта через функцию "
        "calculatePrice (ветка PlisseNetEngine.calculate) без изменения бизнес-логики."
    )
    document.add_paragraph(
        "Для проверки использованы типоразмеры: 600×1200, 800×1600, 1000×2000, "
        "1200×2200, 1400×2400, 1600×2500 мм."
    )
    document.add_paragraph(
        "Итоговая таблица предназначена для заведения номенклатуры в дилерскую оконную программу "
        "по безопасной цене за 1 м²."
    )
    document.add_paragraph(
        "RAL исключен из расчётной части таблиц и не участвует в формировании цены за 1 м²."
    )

    document.add_heading("Итоговая таблица стоимости 1 м² по вариантам", level=2)
    summary_headers = [
        "Группа цвета профиля",
        "Группа полотна",
        "Тип открывания",
        "Мин. цена за 1 м², ₽",
        "Макс. цена за 1 м², ₽",
        "Средняя цена за 1 м², ₽",
        "Рекомендуемая цена за 1 м², ₽",
        "Дилерская цена за 1 м² со скидкой 30%, ₽",
        "Разброс цены за 1 м², %",
        "Рекомендация",
        "Комментарий",
    ]
    summary_table_rows = [
        [
            row["colorGroupLabel"],
            row["meshGroupLabel"],
            row["openingLabel"],
            money(row["minRubPerM2"]),
            money(row["maxRubPerM2"]),
            money(row["avgRubPerM2"]),
            money(row["recommendedRubPerM2"]),
            f"{row['dealerRubPerM2Discount30']:.0f}",
            percent(row["spreadPct"]),
            row["recommendation"],
            row["comment"],
        ]
        for row in summary_rows
    ]
    add_table(document, summary_headers, summary_table_rows)

    if range_rows:
        document.add_heading("Таблица по диапазонам площади", level=2)
        range_headers = [
            "Группа цвета профиля",
            "Группа полотна",
            "Тип открывания",
            "Диапазон площади",
            "Мин. цена за 1 м², ₽",
            "Макс. цена за 1 м², ₽",
            "Средняя цена за 1 м², ₽",
            "Рекомендуемая цена за 1 м², ₽",
            "Кол-во проверок",
        ]
        range_table_rows = [
            [
                row["colorGroupLabel"],
                row["meshGroupLabel"],
                row["openingLabel"],
                row["areaRange"],
                money(row["minRubPerM2"]),
                money(row["maxRubPerM2"]),
                money(row["avgRubPerM2"]),
                money(row["recommendedRubPerM2"]),
                str(row["samples"]),
            ]
            for row in range_rows
        ]
        add_table(document, range_headers, range_table_rows)

    document.add_heading("Проверочная таблица по типоразмерам", level=2)
    checks_headers = [
        "Группа цвета профиля",
        "Группа полотна",
        "Тип открывания",
        "Ширина, мм",
        "Высота, мм",
        "Площадь, м²",
        "Расчетная цена изделия, ₽",
        "Расчетная цена за 1 м², ₽",
        "Дилерская цена изделия со скидкой 30%, ₽",
        "Дилерская цена за 1 м² со скидкой 30%, ₽",
    ]
    checks_table_rows = [
        [
            row["colorGroupLabel"],
            row["meshGroupLabel"],
            row["openingLabel"],
            str(row["width"]),
            str(row["height"]),
            f"{row['areaM2']:.4f}",
            money(row["totalRub"]),
            money(row["rubPerM2"]),
            f"{row['dealerTotalRubDiscount30']:.0f}",
            f"{row['dealerRubPerM2Discount30']:.0f}",
        ]
        for row in check_rows
    ]
    add_table(document, checks_headers, checks_table_rows)

    document.add_heading("Выводы", level=2)
    document.add_paragraph(
        f"1) Единая цена за 1 м²: по {ranges_required_count} из {len(summary_rows)} комбинаций "
        "рекомендуется использовать диапазоны площади."
    )
    document.add_paragraph(
        f"2) Зависимость от размера: разброс цены за 1 м² составляет от {min_spread:.2f}% "
        f"до {max_spread:.2f}%, средний разброс {avg_spread:.2f}%."
    )
    document.add_paragraph(
        "3) Наибольшее влияние на стоимость оказывают тип открывания "
        "(особенно встречное открывание) и группа полотна."
    )
    document.add_paragraph(
        "4) Для практического внедрения в дилерской программе рекомендуется использовать "
        "ставки по диапазонам площади (до 1, 1-2, 2-3, более 3 м²), "
        "так как единая ставка часто приводит к сильному завышению крупных размеров."
    )

    document.add_heading("Комментарии для дилера", level=2)
    document.add_paragraph(
        "Порошковая покраска профиля по RAL рассчитывается отдельно и составляет 200 руб. за погонный метр профиля."
    )

    document.save(OUTPUT_DOCX)
    print(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_doc()
